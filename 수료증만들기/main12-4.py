from docx import Document
from openpyxl import load_workbook
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import glob
load_wb =  load_workbook
(r"C:\Users\user\Downloads\Pyprojectsforme\수료증만들기\수료증명단.xlsx")
# 이 모듈을 이용하여 엑셀 파일을 불러와줌
load_ws =load_wb.active
#워크시트를 선택하여준다 active
name_list = []
birthday_list = []
ho_list = []
for i in range(1,load_ws.max_row + 1): 
    name_list.append(load_ws.cell(i, 1).value)
    birthday_list.append(load_ws.cell(i, 2).value)
    ho_list.append(load_ws.cell(i, 3).value)
#열의 개수만큼을  for문을 돌릴수 있도록 범위를 max_row(최대열 여기에 +1 )
print(name_list)
print(birthday_list)
print(ho_list)

for i in range(len(name_list)):
    doc = docx.Document(r"C:\Users\user\Downloads\Pyprojectsforme\수료증만들기\수료증양식.docx")
    style = doc.styles['Normal']
    style.font.name = '나눔고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(12)
    
    para = doc.add_paragraph()
    run = para.add_run('\n\n') 
    run = para.add_run('              제 '+ ho_list[i] +' 호\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('수  료  증') 
    run.font.name = '나눔고딕'
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(40)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('        성       명: ' + name_list[i] +'\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        생 년 월 일: ' + birthday_list[i] +'\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        교 육 과 정: 파이썬과 40개의 작품들\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20) 
    run = para.add_run('        교 육 날 짜: 2021.08.05~2021.09.09\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('        위 사람은 파이썬과 40개의 작품들 교육과정을\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        이수하였으므로 이 증서를 수여 합니다.\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n\n')
    run = para.add_run('2021.09.19') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n\n')
    run = para.add_run('파이썬교육기관장') 
    run.font.name = '나눔고딕'
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save('수료증만들기\\'+name_list[i]+'.docx')
    #convert를 사용하여 pdf로 변환하여주고 싶었으나....
    #오류가 발생하여 이유를 검색해보았음
    #A 문제 : Pywin32 설치 오류
    #pywintypes.com_error: (-2147221005, '잘못된 클래스 문자열입니다.', None, None)
    #A-1번 이유 :32비트로 동작하기 때문이다. 그러므로 32비트 파이썬을 설치해준다
    # 그후에는 vscode 인터프리티어에서 32비트로 설정후 다시 install 해주어야함
    # -> 오류가 아직 해결되지 않음
    #A-2번 이유 : pywin32가 고장났으므로, 재설치를 하여야한다.
    # -> pip uninstall pywin32를 하여, 지우고 다시 재설치를 해보았음
    #-> 해본 결과, 오류의 해결!
    
    #B 문제 : Docx2pdf 
    #B-1번 이유 : [원인] 관련 프로그램 개발 툴 미설치+ 관리자권한 미부여 
    # 관리자 권한을 부여하여 vscdoe를 실행, 진행하였다
    # -> 결과 오류는 해결되지 않았다.
    #B-2번 이유 : office가 설치되지 아니함 * 
    #-> https://stackoverflow.com/questions/6011115/doc-to-pdf-using-python
    #@Abdelhedihlel Unfortunately, it requires Microsoft Office to be installed and thus only works on Windows and macOS. 
    
    
 
   